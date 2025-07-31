import io
from typing import Optional
import PyPDF2
import docx
import markdown
import chardet

class FileHandler:
    """Handle different file types for citation extraction"""
    
    def extract_text(self, uploaded_file) -> Optional[str]:
        """Extract text from uploaded file"""
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension == 'txt':
                return self._extract_from_txt(uploaded_file)
            elif file_extension == 'pdf':
                return self._extract_from_pdf(uploaded_file)
            elif file_extension == 'docx':
                return self._extract_from_docx(uploaded_file)
            elif file_extension == 'md':
                return self._extract_from_markdown(uploaded_file)
            else:
                return None
                
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            return None
    
    def _extract_from_txt(self, file) -> str:
        """Extract text from TXT file"""
        # Try to detect encoding
        raw_data = file.read()
        file.seek(0)  # Reset file pointer
        
        result = chardet.detect(raw_data)
        encoding = result['encoding'] or 'utf-8'
        
        try:
            return raw_data.decode(encoding)
        except:
            # Fallback to utf-8 with error handling
            return raw_data.decode('utf-8', errors='replace')
    
    def _extract_from_pdf(self, file) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                
                # Clean up text
                page_text = page_text.replace('\n\n', '\n')
                page_text = ' '.join(page_text.split())
                
                text += page_text + "\n\n"
                
        except Exception as e:
            print(f"Error reading PDF: {str(e)}")
            
        return text.strip()
    
    def _extract_from_docx(self, file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file)
            
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            
            # Also check headers and footers
            for section in doc.sections:
                # Header
                header = section.header
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                
                # Footer
                footer = section.footer
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
            
            return '\n\n'.join(full_text)
            
        except Exception as e:
            print(f"Error reading DOCX: {str(e)}")
            return ""
    
    def _extract_from_markdown(self, file) -> str:
        """Extract text from Markdown file"""
        try:
            # Read the markdown content
            content = file.read().decode('utf-8')
            
            # Convert to HTML then extract text
            html = markdown.markdown(content)
            
            # Simple HTML tag removal
            import re
            text = re.sub('<[^<]+?>', '', html)
            
            # Decode HTML entities
            import html as html_module
            text = html_module.unescape(text)
            
            return text.strip()
            
        except Exception as e:
            print(f"Error reading Markdown: {str(e)}")
            return ""
    
    def get_file_info(self, uploaded_file) -> dict:
        """Get information about the uploaded file"""
        return {
            "filename": uploaded_file.name,
            "size": uploaded_file.size,
            "type": uploaded_file.type,
            "extension": uploaded_file.name.split('.')[-1].lower()
        }